""""DOC/DOCX file translation"""
import argparse

import os
import re

import docx

from extension.files.utils import should_translate
from service.mt import translator_factory
from service.utils import detect_lang



def scan_doc(doc):
    runs = []

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if should_translate(run.text):
                runs.append(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if should_translate(run.text):
                            runs.append(run)

    for sec in doc.sections:
        for p in sec.header.paragraphs:
            for run in p.runs:
                if should_translate(run.text):
                    runs.append(run)

        for p in sec.footer.paragraphs:
            for run in p.runs:
                if should_translate(run.text):
                    runs.append(run)

    return runs


def translate_docx_auto(docx_fn, source_lang="auto", target_lang="zh", translation_file=None, callbacker=None):
    paths = os.path.splitext(docx_fn)

    if translation_file is None:
        translated_fn = paths[0] + "-translated.docx"
    else:
        translated_fn = translation_file

    if callbacker:
        callbacker.set_info("读取源文档...", docx_fn)

    doc = docx.Document(docx_fn)  # TODO: 大文档能一次读入？
    runs = scan_doc(doc)

    if source_lang == "auto":
        source_lang = detect_lang(runs[0].text)  # TODO: 语言检测更安全些

    translator = translator_factory.get_translator(source_lang, target_lang)

    if translator is None:
        raise ValueError("给定语言对不支持: {}".format(source_lang+"-"+target_lang))

    txt_list = [r.text for r in runs]
    for t in txt_list:
        print(t)
    result_list = translator.translate_list(txt_list, sl=source_lang, tl=target_lang,
                                            callbacker=callbacker, fn=docx_fn)

    if callbacker:
        callbacker.set_info("翻译完成，写出翻译结果", docx_fn)

    for i in range(len(runs)):
        runs[i].text = result_list[i]

    doc.save(translated_fn)

    return translated_fn


if __name__ == "__main__":
    arg_parser = argparse.ArgumentParser("DOC File Translator")
    arg_parser.add_argument("--to_lang", type=str, default="zh", help="target language")
    arg_parser.add_argument("-i", "--input_file", type=str, required=True, help="file to be translated")
    arg_parser.add_argument("-o", "--output_file", type=str, default=None, help="translation file")
    args = arg_parser.parse_args()

    in_file = args.input_file
    out_file = args.output_file
    to_lang = args.to_lang

    translated_fn = translate_docx_auto(in_file, target_lang=to_lang, translation_file=out_file)

    import webbrowser

    webbrowser.open(in_file)
    webbrowser.open(translated_fn)

